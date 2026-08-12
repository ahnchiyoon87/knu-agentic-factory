"""대표님(1일차 담당)께 보낼 것을 뽑는다.

    python 공유본만들기.py

무엇을 드리나 — **한 문서뿐이다.** PDF(읽기용)와 docx(고쳐 쓰실 때용) 두 형식.

    **2·3일차 실라버스** — 1일차 강의안과 **같은 형식**으로 쓴 것

왜 형식을 맞추나
    대표님 1일차 강의안이 `1. 과정 개요 → 2. 커리큘럼 한눈에 → 3. 모듈별 상세`
    구조이고, 각 모듈이 **체험 키워드 / 목표 / 다루는 내용 / 실습 / 산출물** 다섯 칸이다.
    형식이 다르면 나란히 놓았을 때 따로 논다.

왜 강의안 docx 를 같이 안 넣나
    `강의안/` 의 docx 2개는 **발주처 제출물**이다. 실라버스에 이미 모듈별 상세가
    다 들어 있어 **같은 말을 두 번** 하게 되고, 받는 분이 뭘 먼저 볼지 헷갈린다.
    더 깊은 것이 필요하다고 하시면 그때 `강의안/` 에서 따로 드린다.
"""

from __future__ import annotations

import sys as _sys
for _s in (_sys.stdout, _sys.stderr):
    if (getattr(_s, "encoding", "") or "").lower().replace("-", "") != "utf8":
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

import re
import shutil
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent
REPO = ROOT.parent
SRC = REPO / "제작" / "강의안원본" / "2·3일차_실라버스_공유본.md"
OUT = REPO / "산출물" / "공유"
제목 = "피지컬AI 실습 — 2·3일차 실라버스"


def md2docx(src: Path, dst: Path) -> None:
    """PDF 는 읽기용, docx 는 대표님이 고쳐 쓰실 수 있게."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10.5)

    줄 = src.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(줄):
        line = 줄[i]
        # 그림 — `![설명](경로)`. docx 에도 실제 이미지로 넣는다.
        if m := re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip()):
            from docx.shared import Mm
            그림 = (src.parent / m.group(1 + 1)).resolve()
            if 그림.is_file():
                doc.add_picture(str(그림), width=Mm(160))
                if m.group(1):
                    cap = doc.add_paragraph(m.group(1))
                    for r in cap.runs:
                        r.italic = True; r.font.size = Pt(9)
            i += 1
            continue
        if line.startswith("```"):
            i += 1
            buf = []
            while i < len(줄) and not 줄[i].startswith("```"):
                buf.append(줄[i]); i += 1
            r = doc.add_paragraph().add_run("\n".join(buf))
            r.font.name = "Consolas"; r.font.size = Pt(9.5)
            i += 1
            continue
        if re.match(r"^\s*\|.*\|\s*$", line):
            rows = []
            while i < len(줄) and re.match(r"^\s*\|.*\|\s*$", 줄[i]):
                rows.append([c.strip() for c in 줄[i].strip().strip("|").split("|")]); i += 1
            body = [r for r in rows[1:] if not re.match(r"^[\s:\-]+$", "".join(r))]
            t = doc.add_table(rows=0, cols=len(rows[0])); t.style = "Table Grid"
            cs = t.add_row().cells
            for j, c in enumerate(rows[0]):
                cs[j].text = re.sub(r"[*`]", "", c)
                for pp in cs[j].paragraphs:
                    for rr in pp.runs: rr.bold = True
            for row in body:
                cc = t.add_row().cells
                for j, c in enumerate(row[:len(rows[0])]):
                    cc[j].text = re.sub(r"[*`]", "", c)
            doc.add_paragraph()
            continue
        if m := re.match(r"^(#{1,3})\s+(.*)$", line):
            h = doc.add_heading(re.sub(r"[*`]", "", m.group(2)), level=len(m.group(1)))
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x14, 0x33, 0x5E); r.font.name = "맑은 고딕"
        elif line.startswith(">"):
            buf = []
            while i < len(줄) and 줄[i].startswith(">"):
                buf.append(줄[i].lstrip(">").strip()); i += 1
            for x in [b for b in buf if b]:
                p = doc.add_paragraph(re.sub(r"[*`]", "", x))
                p.paragraph_format.left_indent = Pt(18)
                for r in p.runs: r.italic = True
            continue
        elif re.match(r"^\s*[-*]\s+", line):
            doc.add_paragraph(re.sub(r"[*`]", "", re.sub(r"^\s*[-*]\s+", "", line)),
                              style="List Bullet")
        elif re.match(r"^\s*\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"[*`]", "", re.sub(r"^\s*\d+\.\s+", "", line)),
                              style="List Number")
        elif re.match(r"^\s*---+\s*$", line):
            doc.add_paragraph()
        elif line.strip():
            doc.add_paragraph(re.sub(r"[*`]", "", line))
        i += 1
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


def main() -> int:
    if not SRC.is_file():
        print(f"원본이 없습니다: {SRC}", file=sys.stderr)
        return 1
    if OUT.exists():
        shutil.rmtree(OUT)
    OUT.mkdir(parents=True)

    pdf = OUT / f"{제목}.pdf"
    r = subprocess.run([sys.executable, str(ROOT / "md2pdf.py"), str(SRC), str(pdf), 제목],
                       capture_output=True, text=True, encoding="utf-8", errors="replace")
    if r.returncode != 0 or not pdf.is_file():
        print(f"PDF 변환 실패\n{r.stdout}{r.stderr}", file=sys.stderr)
        return 1
    print(f"  만듦  {pdf.name}  ({pdf.stat().st_size // 1024}KB)")

    docx = OUT / f"{제목}.docx"
    md2docx(SRC, docx)
    print(f"  만듦  {docx.name}  ({docx.stat().st_size // 1024}KB)")

    print(f"\n  {OUT}")
    print("  이 폴더를 통째로 보내시면 됩니다. (같은 문서의 PDF · docx 두 형식)")
    print("\n  ※ 더 깊은 것을 원하시면 그때 `강의안/` 의 docx 를 따로 드립니다.")
    print("     실라버스에 이미 모듈별 상세가 들어 있어 같이 주면 같은 말이 두 번 됩니다.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
